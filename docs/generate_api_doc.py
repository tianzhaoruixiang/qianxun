#!/usr/bin/env python3
"""生成千寻 Claude Code SDK 对外接口文档 (DOCX)。"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "QianXun-ClaudeCode-SDK-API-文档.docx")


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = bold
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_api_table(doc, rows):
    """rows: list of (method, path, auth, desc)"""
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["方法", "路径", "鉴权", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "2B579A")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9)
    for ri, (method, path, auth, desc) in enumerate(rows, start=1):
        table.rows[ri].cells[0].text = method
        table.rows[ri].cells[1].text = path
        table.rows[ri].cells[2].text = auth
        table.rows[ri].cells[3].text = desc
        for cell in table.rows[ri].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def build():
    doc = Document()
    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("千寻智能体平台\nClaude Code SDK 对外接口文档")
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.name = "Microsoft YaHei"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"版本 1.0  |  生成日期 {date.today().isoformat()}")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # 目录说明
    add_heading(doc, "文档说明", 1)
    add_para(doc, "本文档描述千寻（QianXun）后端对外 REST API，用于集成 Claude Code Agent SDK 网关，"
                   "覆盖多智能体协作、智能体管理、运行时可观测性、技能、工具集及插件扩展等能力。")
    add_para(doc, "基础 URL：http(s)://{host}:8080")
    add_para(doc, "所有业务接口前缀：/QianXunService")
    add_para(doc, "Claude Code 网关（内部/运维直连，非前端暴露）：{qianxun.claude.base-url}，默认 http://claude-code:8642")

    add_heading(doc, "1. 通用约定", 1)

    add_heading(doc, "1.1 统一响应格式", 2)
    add_code(doc, '{\n  "code": 0,\n  "message": "success",\n  "data": <T>\n}')
    add_para(doc, "code=0 表示成功；业务错误返回非零 code 与 message；401 由 JWT 过滤器直接返回，非 ApiResponse 包装。")

    add_heading(doc, "1.2 请求包装（部分 POST 接口）", 2)
    add_code(doc, '{\n  "jsonArg": { /* 业务参数 */ },\n  "generalArgument": {\n    "userId": "可选，覆盖当前用户上下文",\n    "loginName": "",\n    "ip": "",\n    "coralKey": ""\n  }\n}')
    add_para(doc, "Hermes 管理面、Registry、Session 等 POST 接口使用 ApiRequest 包装；Chat Stream 直接使用 StreamChatRequest。")

    add_heading(doc, "1.3 鉴权", 2)
    add_para(doc, "推荐：Authorization: Bearer <JWT>（POST /QianXunService/auth/login 获取）")
    add_para(doc, "开发回退：X-User-Id（必填）、X-User-Name、X-User-Display-Name（URL 编码）")
    add_para(doc, "管理员（role=admin）：智能体注册表 upsert/delete、Profile 创建、Soul 更新")
    add_para(doc, "公开白名单：GET /auth/health、POST /auth/login、GET /data/files/public/**")

    add_heading(doc, "1.4 架构概览", 2)
    add_code(doc, """前端/第三方
    ↓ REST + SSE
千寻后端 (Spring Boot, :8080)
    ├─ HermesAgentClient → Claude 网关管理面
    │     GET/POST/PUT/DELETE /api/profiles|skills|tools|config|files|mcp|plugins|delegation
    └─ ClaudeCodeChatClient → POST /v1/agent/stream (NDJSON 流式对话)

用户数据目录：
  /opt/data/{userId}/profiles/{profile}/   — CLAUDE.md、.claude/skills
  /opt/data/{userId}/workspace/              — 智能体工作区 cwd
  cache/delegation/live/deleg_*/             — 子智能体 live transcript""")

    doc.add_page_break()

    # 2. 认证
    add_heading(doc, "2. 认证接口", 1)
    add_api_table(doc, [
        ("GET", "/QianXunService/auth/health", "公开", "健康检查"),
        ("POST", "/QianXunService/auth/login", "公开", "用户名密码登录，返回 JWT"),
    ])
    add_heading(doc, "2.1 POST /auth/login", 2)
    add_para(doc, "请求体：", bold=True)
    add_code(doc, '{ "username": "admin", "password": "admin123" }')
    add_para(doc, "响应 data：LoginResponse", bold=True)
    add_code(doc, '{\n  "token": "eyJ...",\n  "expiresInSeconds": 86400,\n  "userId": "1",\n  "username": "admin",\n  "displayName": "管理员",\n  "role": "admin"\n}')

    doc.add_page_break()

    # 3. 智能体管理
    add_heading(doc, "3. 智能体管理", 1)
    add_para(doc, "智能体由两层组成：")
    add_para(doc, "• agent_registry（千寻注册表）：面向用户的智能体超市元数据")
    add_para(doc, "• Claude Profile（Hermes Profile）：运行时身份，含 CLAUDE.md（Soul）、技能、工具集配置")

    add_heading(doc, "3.1 智能体注册表 — /registry", 2)
    add_api_table(doc, [
        ("POST", "/registry/agents/list", "登录", "列出注册智能体；jsonArg.enabledOnly 默认 true"),
        ("POST", "/registry/agents/upsert", "管理员", "创建/更新智能体，同步创建 Profile + Soul + 发布模板"),
        ("POST", "/registry/agents/delete", "管理员", "删除智能体及关联 Claude Profile"),
        ("POST", "/registry/models/list", "登录", "模型注册表列表"),
        ("POST", "/registry/models/upsert", "登录", "模型 upsert"),
        ("POST", "/registry/datasets/list", "登录", "数据集注册表列表"),
        ("POST", "/registry/datasets/upsert", "登录", "数据集 upsert"),
    ])

    add_heading(doc, "3.1.1 UpsertAgentRegistryRequest", 3)
    add_code(doc, """{
  "code": "analyst",           // 必填，唯一标识
  "name": "数据分析师",         // 必填
  "soulMd": "# 角色定义...",   // 必填，写入 CLAUDE.md
  "category": "general",
  "description": "",
  "icon": "",
  "modelCode": "",
  "welcomeTitle": "",
  "welcomeIntro": "",
  "presetChat1": "",
  "presetChat2": "",
  "presetChat3": "",
  "hermesProfile": "",         // 可选，默认用 code
  "priority": 100,
  "enabled": true
}""")

    add_heading(doc, "3.1.2 AgentRegistryResponse", 3)
    add_code(doc, """{
  "id": "...",
  "code": "analyst",
  "name": "数据分析师",
  "category": "general",
  "description": "",
  "icon": "",
  "modelCode": "",
  "welcomeTitle": "",
  "welcomeIntro": "",
  "presetChat1": "",
  "presetChat2": "",
  "presetChat3": "",
  "hermesProfile": "analyst",
  "priority": 100,
  "enabled": true
}""")

    add_heading(doc, "3.2 Profile / Soul 管理 — /hermes/profiles", 2)
    add_api_table(doc, [
        ("POST", "/hermes/profiles/list", "登录", "列出当前用户的 Claude Profiles"),
        ("POST", "/hermes/profiles/create", "管理员", "创建新 Profile"),
        ("POST", "/hermes/profiles/soul", "登录", "读取 CLAUDE.md（Soul）"),
        ("POST", "/hermes/profiles/soul/update", "管理员", "写入 CLAUDE.md"),
    ])

    add_heading(doc, "3.2.1 HermesProfileResponse", 3)
    add_code(doc, '{ "name": "analyst", "description": "", "model": "", "active": true, "path": "...", "contextWindow": 200000 }')

    add_heading(doc, "3.2.2 CreateHermesProfileRequest", 3)
    add_code(doc, '{ "name": "my-agent", "description": "自定义智能体" }')

    add_heading(doc, "3.2.3 HermesSoulRequest / HermesSoulResponse", 3)
    add_code(doc, '请求: { "name": "analyst", "content": "# Soul 内容" }\n响应: { "name": "analyst", "content": "...", "exists": true }')

    add_heading(doc, "3.3 会话管理 — /sessions", 2)
    add_api_table(doc, [
        ("POST", "/sessions/create", "登录", "创建对话会话"),
        ("POST", "/sessions/list", "登录", "分页列表"),
        ("POST", "/sessions/get", "登录", "单条详情"),
        ("POST", "/sessions/update", "登录", "更新标题/长程目标"),
        ("POST", "/sessions/delete", "登录", "删除会话"),
        ("POST", "/sessions/messages", "登录", "消息历史"),
    ])

    add_heading(doc, "3.3.1 CreateSessionRequest", 3)
    add_code(doc, '{ "title": "新对话", "agentCode": "analyst", "hermesProfile": "analyst", "agentName": "数据分析师" }')

    add_heading(doc, "3.3.2 ChatSessionResponse", 3)
    add_code(doc, """{
  "id": "session-uuid",
  "title": "新对话",
  "createdAt": "2026-08-22T...",
  "updatedAt": "...",
  "messageCount": 10,
  "lastMessagePreview": "...",
  "agentCode": "analyst",
  "hermesProfile": "analyst",
  "agentName": "数据分析师",
  "goal": { "title": "", "description": "", "steps": "", "constraints": "" },
  "streaming": false
}""")

    doc.add_page_break()

    # 4. 流式对话与多智能体协作
    add_heading(doc, "4. 流式对话与多智能体协作", 1)

    add_heading(doc, "4.1 流式对话 — /sessions/{sessionId}/chat", 2)
    add_api_table(doc, [
        ("POST", "/sessions/{sessionId}/chat/stream", "登录+会话归属", "发起 Claude Code 流式对话（SSE）"),
        ("GET", "/sessions/{sessionId}/chat/stream/subscribe?afterSeq=0", "登录", "断线重连订阅 SSE"),
        ("GET", "/sessions/{sessionId}/chat/runs/active", "登录", "查询进行中 Run（无则 204）"),
        ("POST", "/sessions/{sessionId}/chat/stream/stop", "登录", "取消当前 Run"),
    ])

    add_heading(doc, "4.1.1 StreamChatRequest", 3)
    add_code(doc, """{
  "content": "用户消息文本",
  "modelCode": "",              // 可选，模型注册表 code
  "agentCode": "analyst",       // 可选，智能体 code
  "hermesProfile": "analyst",   // 可选，Claude Profile
  "fileIds": ["file-id-1"],     // 可选，本轮附件 data_file.id
  "skillName": "code-review",   // 可选，强制技能
  "goal": {                     // 可选，设定长程目标
    "title": "完成报告",
    "description": "...",
    "steps": "...",
    "constraints": "..."
  },
  "clearGoal": false,           // true 清除长程目标
  "agentsStatus": false         // true 查询子智能体/委派任务状态
}""")

    add_heading(doc, "4.1.2 SSE 事件类型", 3)
    events = [
        ("started", "流开始，含 sessionId"),
        ("session_goal", "长程目标状态 { cleared, goal }"),
        ("token", "增量文本 { text }"),
        ("tool_call", "工具调用 { id, name, displayName, status, input, output, ... }"),
        ("usage", "Token 用量 { promptTokens, completionTokens, totalTokens }"),
        ("generated_file", "智能体生成的文档"),
        ("suggestions", "后续建议问题列表"),
        ("stream_warning", "非致命警告"),
        ("done", "正常结束 { finishReason, usage }"),
        ("error", "错误信息"),
    ]
    t = doc.add_table(rows=1 + len(events), cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "事件名"
    t.rows[0].cells[1].text = "说明"
    set_cell_shading(t.rows[0].cells[0], "2B579A")
    set_cell_shading(t.rows[0].cells[1], "2B579A")
    for i, (ev, desc) in enumerate(events, 1):
        t.rows[i].cells[0].text = ev
        t.rows[i].cells[1].text = desc
    doc.add_paragraph()

    add_heading(doc, "4.1.3 ActiveRunResponse", 3)
    add_code(doc, '{ "runId": "...", "sessionId": "...", "status": "RUNNING", "assistantMessageId": "...", "lastSeq": 42, "cancelRequested": null }')

    add_heading(doc, "4.1.4 并发控制", 3)
    add_para(doc, "同一会话同时只能有一个 Run；冲突时 POST /stream 返回 HTTP 409：")
    add_code(doc, '{ "code": 409, "message": "该会话正在输出中，请等待完成或先停止后再发送", "data": null }')

    add_heading(doc, "4.2 多智能体协作机制", 2)
    add_para(doc, "多智能体协作通过 Claude Code 内置 delegation 工具集实现，非独立 REST 调度：")
    add_para(doc, "• 工具集 delegation 包含：Agent、Task、SendMessage")
    add_para(doc, "• 用户可通过 agentsStatus=true 或斜杠命令 /agents、/tasks、/task 查询子智能体状态")
    add_para(doc, "• 委派执行日志落盘至 cache/delegation/live/deleg_*/task-N.log")
    add_para(doc, "• ClaudeCodeChatClient 将 /goal、/agents、/skill 等斜杠命令改写为自然语言 prompt 下发网关")

    add_heading(doc, "4.3 委派观测 — /hermes/delegation/live", 2)
    add_api_table(doc, [
        ("POST", "/hermes/delegation/live/list", "登录", "列出最近委派 transcript"),
        ("POST", "/hermes/delegation/live/read", "登录", "读取指定委派日志内容"),
    ])

    add_heading(doc, "4.3.1 HermesLiveTranscriptListRequest", 3)
    add_code(doc, '{ "profile": "analyst", "limit": 20 }')

    add_heading(doc, "4.3.2 HermesLiveDelegationResponse", 3)
    add_code(doc, """{
  "delegationId": "deleg_abc123",
  "path": "cache/delegation/live/deleg_abc123",
  "started": "2026-08-22T...",
  "completed": true,
  "taskCount": 2,
  "tasks": [
    { "index": 0, "path": ".../task-0.log", "goal": "...", "status": "completed", "size": 4096 }
  ]
}""")

    add_heading(doc, "4.3.3 HermesLiveTranscriptReadRequest", 3)
    add_code(doc, '{ "profile": "analyst", "delegationId": "deleg_abc123", "taskIndex": 0, "maxChars": 50000 }')

    doc.add_page_break()

    # 5. 技能管理
    add_heading(doc, "5. 技能（Skills）管理 — /hermes/skills", 1)
    add_api_table(doc, [
        ("POST", "/hermes/skills/list", "登录", "列出 Profile 下技能"),
        ("POST", "/hermes/skills/tree", "登录", "技能文件树"),
        ("POST", "/hermes/skills/file", "登录", "读取技能文件（默认 SKILL.md）"),
        ("POST", "/hermes/skills/file/update", "登录", "写入技能文件"),
        ("POST", "/hermes/skills/toggle", "登录", "启用/禁用技能"),
        ("POST", "/hermes/skills/upload", "登录", "multipart 上传 .zip 技能包"),
        ("GET", "/hermes/skills/download?name=&profile=", "登录", "下载技能 zip"),
    ])

    add_heading(doc, "5.1 HermesSkillItemResponse", 2)
    add_code(doc, '{ "name": "code-review", "description": "...", "category": "dev", "enabled": true, "provenance": "upload" }')

    add_heading(doc, "5.2 HermesSkillFileRequest", 2)
    add_code(doc, '{ "profile": "analyst", "name": "code-review", "path": "SKILL.md", "content": "...", "enabled": true }')

    add_heading(doc, "5.3 技能触发方式", 2)
    add_para(doc, "1. StreamChatRequest.skillName — 对话中强制使用指定技能")
    add_para(doc, "2. 用户输入 /{skill-slug} — Dashboard 原生斜杠命令")
    add_para(doc, "技能存储路径：/opt/data/{userId}/profiles/{profile}/.claude/skills/{name}/")

    doc.add_page_break()

    # 6. 工具集管理
    add_heading(doc, "6. 工具集（Tools）管理 — /hermes/tools", 1)
    add_api_table(doc, [
        ("POST", "/hermes/tools/list", "登录", "列出工具集及原子工具"),
        ("POST", "/hermes/tools/toggle", "登录", "开关工具集"),
    ])

    add_heading(doc, "6.1 内置工具集目录", 2)
    toolsets = [
        ("web", "Web", "WebSearch, WebFetch", "默认启用"),
        ("file", "File", "Read, Write, Edit, Glob, Grep, NotebookEdit", "默认启用"),
        ("terminal", "Terminal", "Bash", "默认启用"),
        ("code_execution", "Code", "Bash", "默认启用"),
        ("delegation", "Delegation", "Agent, Task, SendMessage", "默认启用，多智能体"),
        ("skills", "Skills", "Skill", "默认关闭"),
        ("todo", "Todo", "TodoWrite, TaskCreate, TaskGet, TaskList, TaskUpdate", "默认关闭"),
        ("memory", "Memory", "—", "默认关闭"),
        ("session_search", "Session search", "—", "默认关闭"),
        ("browser", "Browser", "—", "默认关闭"),
        ("kanban", "Kanban", "TodoWrite, TaskCreate, TaskList, TaskUpdate", "默认关闭"),
    ]
    tt = doc.add_table(rows=1 + len(toolsets), cols=4)
    tt.style = "Table Grid"
    for i, h in enumerate(["名称", "标签", "Claude 工具", "默认状态"]):
        tt.rows[0].cells[i].text = h
        set_cell_shading(tt.rows[0].cells[i], "2B579A")
    for ri, row in enumerate(toolsets, 1):
        for ci, val in enumerate(row):
            tt.rows[ri].cells[ci].text = val
    doc.add_paragraph()

    add_heading(doc, "6.2 HermesToolsetItemResponse", 2)
    add_code(doc, """{
  "name": "web",
  "label": "Web",
  "description": "网页搜索与抓取",
  "platform": "cli",
  "platformLabel": "CLI",
  "enabled": true,
  "configured": true,
  "tools": [
    { "name": "WebSearch", "displayName": "网页搜索", "iconKind": "search", "enabled": true }
  ]
}""")

    add_heading(doc, "6.3 HermesToolsetToggleRequest", 2)
    add_code(doc, '{ "profile": "analyst", "name": "delegation", "enabled": true }')

    add_heading(doc, "6.4 对话网关工具集同步", 2)
    add_para(doc, "HermesToolsetService 在 toggle 后同步 Claude 网关 GET/PUT /api/config，"
                   "ClaudeCodeChatClient 流式对话时读取 enabled 列表作为 allowedToolsets。")
    add_para(doc, "MCP 运行时：ClaudeCodeChatClient 在 isChatMcpAllowed 为 true 时不传 mcpDisabled；"
                   "chat.js 加载 profile 下已启用的 MCP Server 并注入 SDK mcpServers。"
                   "no_mcp 哨兵或 qianxun.claude.append-no-mcp=true 时对话禁用 MCP。")

    add_heading(doc, "6.5 MCP 与插件 — /hermes/mcp、/hermes/plugins", 2)
    add_api_table(doc, [
        ("POST", "/hermes/mcp/list", "登录", "列出 MCP Server"),
        ("POST", "/hermes/mcp/upsert", "登录", "创建/更新 MCP Server"),
        ("POST", "/hermes/mcp/toggle", "登录", "启用/禁用 MCP Server"),
        ("POST", "/hermes/mcp/delete", "登录", "删除 MCP Server"),
        ("POST", "/hermes/plugins/list", "登录", "列出插件 manifest"),
        ("POST", "/hermes/plugins/upsert", "登录", "创建/更新插件"),
        ("POST", "/hermes/plugins/toggle", "登录", "启用/禁用插件"),
        ("POST", "/hermes/plugins/delete", "登录", "删除插件"),
        ("GET", "/hermes/gateway/status", "登录", "代理 Claude 网关 /api/status"),
    ])
    add_code(doc, '{ "profile": "analyst", "name": "filesystem", "command": "npx", "enabled": true }')

    doc.add_page_break()

    # 7. Claude 网关内部 API
    add_heading(doc, "7. Claude Code 网关 API（参考）", 1)
    add_para(doc, "以下接口由千寻后端 HermesAgentClient / ClaudeCodeChatClient 内部调用。"
                   "第三方集成建议优先使用第 3–6 节千寻 REST 封装；如需直连网关，需配置 CLAUDE_GATEWAY_KEY。")

    add_api_table(doc, [
        ("GET", "/health", "公开", "健康检查"),
        ("GET", "/api/status", "Bearer", "网关状态（脱敏，不含 dataDir）"),
        ("GET", "/api/profiles?userId=", "Bearer", "列出 profiles"),
        ("POST", "/api/profiles", "Bearer", "创建 profile { name, description, userId }"),
        ("GET", "/api/profiles/{name}/soul?userId=", "Bearer", "读取 CLAUDE.md"),
        ("PUT", "/api/profiles/{name}/soul?userId=", "Bearer", "写入 CLAUDE.md"),
        ("POST", "/api/profiles/{name}/publish-template?userId=", "Bearer", "发布模板到 _templates"),
        ("DELETE", "/api/profiles/{name}?userId=", "Bearer", "删除 profile"),
        ("GET", "/api/skills?userId=&profile=", "Bearer", "技能列表"),
        ("GET", "/api/skills/content?name=", "Bearer", "读 SKILL.md"),
        ("PUT", "/api/skills/content", "Bearer", "写 SKILL.md"),
        ("POST", "/api/skills", "Bearer", "创建技能"),
        ("PUT", "/api/skills/toggle", "Bearer", "启用/禁用技能"),
        ("GET", "/api/tools/toolsets?userId=&profile=", "Bearer", "工具集列表"),
        ("PUT", "/api/tools/toolsets/{name}", "Bearer", "开关工具集"),
        ("GET/PUT", "/api/config?userId=&profile=", "Bearer", "对话网关工具集配置"),
        ("GET", "/api/mcp?userId=&profile=", "Bearer", "MCP Server 列表"),
        ("POST", "/api/mcp", "Bearer", "创建/更新 MCP Server"),
        ("PUT", "/api/mcp/{name}/toggle", "Bearer", "启用/禁用 MCP"),
        ("DELETE", "/api/mcp/{name}", "Bearer", "删除 MCP Server"),
        ("GET", "/api/plugins?userId=&profile=", "Bearer", "插件列表"),
        ("POST", "/api/plugins", "Bearer", "创建/更新插件"),
        ("PUT", "/api/plugins/{name}/toggle", "Bearer", "启用/禁用插件"),
        ("DELETE", "/api/plugins/{name}", "Bearer", "删除插件"),
        ("POST", "/api/delegation/cancel", "Bearer", "向运行中委派发送取消信号"),
        ("GET", "/api/files?path=", "Bearer", "列目录"),
        ("POST", "/api/files/mkdir", "Bearer", "建目录"),
        ("POST", "/api/files/write", "Bearer", "写文件（base64）"),
        ("GET", "/api/files/download?path=", "Bearer", "下载文件"),
        ("DELETE", "/api/files?path=", "Bearer", "删除文件/目录"),
        ("POST", "/v1/agent/stream", "Bearer", "NDJSON 流式对话（Accept: application/x-ndjson）"),
    ])

    add_heading(doc, "7.1 POST /v1/agent/stream 请求体", 2)
    add_code(doc, """{
  "sessionId": "会话缓存键",
  "userId": "用户 ID",
  "profile": "analyst",
  "prompt": "用户 prompt（含斜杠命令改写）",
  "seedHistory": [ /* 可选历史 */ ],
  "model": "qwen3.6-plus",
  "permissionMode": "bypassPermissions",
  "allowedToolsets": ["web", "file", "terminal", "delegation"],
  "mcpDisabled": false
}""")

    add_heading(doc, "7.2 NDJSON 流事件（ClaudeCodeStreamParser）", 2)
    add_para(doc, "网关返回 application/x-ndjson，每行一个 JSON 事件：stream_event、assistant、user、result、error。"
                   "解析后映射为 token、tool_call、usage 供 SSE 推送。")

    doc.add_page_break()

    # 8. 运行时可观测性
    add_heading(doc, "8. 运行时可观测性", 1)
    add_api_table(doc, [
        ("GET", "/sessions/{id}/chat/runs/active", "登录", "单会话活跃 Run"),
        ("GET", "/sessions/{id}/chat/stream/subscribe", "登录", "SSE 重连（afterSeq 断点续传）"),
        ("POST", "/runs/list", "登录", "跨会话 Run 列表"),
        ("GET", "/runs/metrics", "登录", "Run 指标快照"),
        ("GET", "/prometheus", "登录", "Prometheus 指标"),
        ("POST", "/hermes/delegation/live/list", "登录", "委派 transcript 列表"),
        ("POST", "/hermes/delegation/live/read", "登录", "委派日志全文/截断"),
        ("POST", "/hermes/delegation/get", "登录", "单个委派详情"),
        ("POST", "/hermes/delegation/cancel", "登录", "取消委派（manifest + 网关 abort 信号）"),
        ("POST", "/hermes/delegation/delete", "登录", "删除委派目录"),
        ("POST", "/activity-logs/list", "登录", "LLM 调用活动日志（开发/运营）"),
    ])
    add_para(doc, "SSE tool_call 事件实时推送工具调用状态；ChatMessage 持久化 toolCallsJson、usageJson。"
                   "Spring Actuator 端点可通过 management 配置暴露。")

    doc.add_page_break()

    # 9. 插件管理
    add_heading(doc, "9. 插件（Plugin）管理 — /hermes/plugins", 1)
    add_para(doc, "插件 manifest 通过 Hermes REST 管理，网关 syncPluginsManifest 同步至 profile 目录；"
                   "已启用插件名称写入对话 system append。")
    add_para(doc, "• POST /hermes/plugins/list|upsert|toggle|delete — 与 MCP 同级的 manifest CRUD")
    add_para(doc, "• MCP Server 配置见第 6.5 节；对话运行时 chat.js 加载 enabled MCP 并注入 SDK")
    add_para(doc, "• no_mcp 工具集哨兵或 append-no-mcp 配置可全局禁用 MCP 运行时")

    doc.add_page_break()

    # 10. 待完善项
    add_heading(doc, "10. 能力缺口与完善建议", 1)

    gaps = [
        ("多智能体协作", "高", [
            "缺少 REST 创建/调度/停止子智能体的显式 API（目前仅通过对话 delegation 工具隐式委派）",
            "无 TeamCreate/TeamDelete 等团队编排管理面",
            "无跨会话委派状态聚合、委派超时/优先级配置",
            "Live transcript 无 WebSocket 实时推送，需轮询 list/read",
        ]),
        ("智能体管理", "高", [
            "Profile 无独立 update/delete REST（delete 仅在 registry/agents/delete 中间接触发）",
            "无运行时生命周期 API（pause/resume/restart agent run）",
            "无智能体版本管理、A/B 测试、灰度发布",
            "无 per-agent 独立模型/工具策略动态调整 REST（仅 upsert 时一次性绑定）",
            "缺少 Profile 列表分页、搜索、标签分类",
        ]),
        ("运行时可观测性", "中", [
            "跨会话 Run 列表与 metrics 已提供；无 OpenTelemetry 专用 trace",
            "Prometheus 端点 /prometheus 已暴露",
            "无结构化 trace ID 贯穿 SSE → 网关 → 子智能体",
            "ActivityLog 面向开发自检，缺少面向运维的告警/仪表盘集成",
            "Token 用量无按用户/智能体/时间聚合统计 API",
        ]),
        ("技能（Skills）", "中", [
            "无技能市场/跨用户共享/评分机制",
            "无技能版本管理与依赖声明",
            "无技能调用统计（次数、成功率、耗时）",
            "无独立 REST 触发单次技能执行（仅 chat skillName 或斜杠命令）",
            "createSkill 网关 API 未暴露为千寻 REST（仅 file/update 间接创建）",
        ]),
        ("工具（Tools）", "中", [
            "仅支持工具集粒度开关，无单工具级 toggle",
            "MCP Server 注册/配置 REST 已提供；健康检查与运行时探针仍待完善",
            "browser/memory/session_search 等 catalog 存在但无配置 REST",
            "无自定义 Function/Tool 注册接口",
            "工具调用审计日志无独立查询 API",
        ]),
        ("插件（Plugins）", "中", [
            "manifest CRUD 与 toggle 已提供；无独立安装包上传/依赖解析",
            "无 plugin 权限沙箱策略 API",
            "无插件依赖管理与冲突检测",
            "MCP vs Plugin vs Skill 边界已在文档中区分，统一注册模型仍待完善",
        ]),
        ("API 工程化", "中", [
            "SpringDoc OpenAPI 已覆盖主要 Controller；可导出 swagger-ui",
            "POST 为主，部分只读接口可补充 GET 语义化路径",
            "无 API 版本前缀（/v1/）",
            "Webhook/回调机制缺失（Run 完成、委派结束通知）",
            "Rate limiting 与配额管理未文档化",
        ]),
    ]

    for area, priority, items in gaps:
        add_heading(doc, f"10.{gaps.index((area, priority, items)) + 1} {area}（优先级：{priority}）", 2)
        for item in items:
            p = doc.add_paragraph(item, style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_page_break()

    # 11. 配置参考
    add_heading(doc, "11. 环境配置参考", 1)
    add_code(doc, """# application.yml / 环境变量
qianxun.claude.enabled: QIANXUN_CLAUDE_ENABLED (默认 true)
qianxun.claude.base-url: QIANXUN_CLAUDE_BASE_URL (默认 http://claude-code:8642)
qianxun.claude.api-key: CLAUDE_GATEWAY_KEY
qianxun.claude.chat-model: QIANXUN_CLAUDE_MODEL (默认 qwen3.6-plus)
qianxun.claude.permission-mode: QIANXUN_CLAUDE_PERMISSION_MODE (默认 bypassPermissions)
qianxun.claude.append-no-mcp: QIANXUN_CLAUDE_APPEND_NO_MCP (默认 false)
qianxun.auth.enabled: QIANXUN_AUTH_ENABLED (默认 true)
qianxun.auth.jwt-secret: QIANXUN_JWT_SECRET

# Claude 网关（docker/claudecode）
CLAUDE_GATEWAY_KEY: 必填 Bearer 鉴权（生产）
CLAUDE_GATEWAY_ALLOW_INSECURE=true: 本地调试免鉴权（勿用于生产）""")

    add_heading(doc, "12. 典型集成流程", 1)
    steps = [
        "POST /auth/login 获取 JWT",
        "POST /registry/agents/list 获取可用智能体",
        "POST /sessions/create 创建会话（携带 agentCode / hermesProfile）",
        "POST /hermes/tools/list 查看/配置工具集（可选）",
        "POST /hermes/mcp/list 配置 MCP Server（可选）",
        "POST /hermes/skills/list 查看可用技能（可选）",
        "POST /sessions/{id}/chat/stream 发起对话（SSE 消费 token/tool_call/done）",
        "GET /sessions/{id}/chat/stream/subscribe 断线重连",
        "POST /hermes/delegation/live/list 观测子智能体委派（多智能体场景）",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}")

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build()
